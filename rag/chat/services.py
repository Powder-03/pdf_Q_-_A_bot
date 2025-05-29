import os
import pickle
from typing import List, Optional
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_community.vectorstores import FAISS
from langchain_core.runnables import RunnableParallel, RunnableLambda, RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.prompts import PromptTemplate
from langchain_core.documents import Document as LangchainDocument
from docx import Document as DocxDocument
from django.conf import settings
from .models import Document, DocumentChunk

class DocumentProcessor:
    def __init__(self):
        self.embeddings = GoogleGenerativeAIEmbeddings(model="models/text-embedding-004")
        self.model = ChatGoogleGenerativeAI(model='gemini-2.5-flash-preview-04-17', temperature=0.7)
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )
        self.prompt = PromptTemplate(
            input_variables=["context", "question"],
            template="""Use the following pieces of context to answer the question at the end. 
If you don't know the answer based on the context, just say that the information is not available in the provided context.

Context:
{context}

Question: {question}

Answer:"""
        )
    
    def extract_text_from_file(self, file_path: str, file_type: str) -> str:
        """Extract text content from different file types"""
        try:
            if file_type.lower() == 'pdf':
                loader = PyPDFLoader(file_path)
                pages = loader.load()
                return "\n\n".join([page.page_content for page in pages])
            
            elif file_type.lower() == 'txt':
                loader = TextLoader(file_path)
                documents = loader.load()
                return "\n\n".join([doc.page_content for doc in documents])
            
            elif file_type.lower() == 'docx':
                doc = DocxDocument(file_path)
                paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
                return "\n\n".join(paragraphs)
            
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            raise Exception(f"Error extracting text from {file_type} file: {str(e)}")
    
    def process_document(self, document: Document) -> bool:
        """Process document and create vector store"""
        try:
            # Extract text content
            file_path = document.file.path
            content = self.extract_text_from_file(file_path, document.file_type)
            
            # Save content to document
            document.content = content
            document.save()
            
            # Create chunks
            langchain_docs = [LangchainDocument(page_content=content)]
            chunks = self.text_splitter.split_documents(langchain_docs)
            
            # Save chunks to database
            DocumentChunk.objects.filter(document=document).delete()
            for i, chunk in enumerate(chunks):
                DocumentChunk.objects.create(
                    document=document,
                    content=chunk.page_content,
                    chunk_index=i
                )
            
            # Create and save vector store
            vectorstore = FAISS.from_documents(chunks, self.embeddings)
            
            # Save vectorstore to file
            vectorstore_path = self.get_vectorstore_path(document.id)
            vectorstore.save_local(vectorstore_path)
            
            # Mark as processed
            document.processed = True
            document.save()
            
            return True
            
        except Exception as e:
            print(f"Error processing document {document.id}: {str(e)}")
            return False
    
    def get_vectorstore_path(self, document_id: str) -> str:
        """Get path for storing vectorstore"""
        vectorstore_dir = os.path.join(settings.MEDIA_ROOT, 'vectorstores')
        os.makedirs(vectorstore_dir, exist_ok=True)
        return os.path.join(vectorstore_dir, str(document_id))
    
    def load_vectorstore(self, document_id: str) -> Optional[FAISS]:
        """Load vectorstore from file"""
        try:
            vectorstore_path = self.get_vectorstore_path(document_id)
            return FAISS.load_local(vectorstore_path, self.embeddings, allow_dangerous_deserialization=True)
        except Exception as e:
            print(f"Error loading vectorstore for document {document_id}: {str(e)}")
            return None
    
    def ask_question(self, document_id: str, question: str, search_kwargs: dict = None) -> str:
        """Ask question about document using RAG"""
        try:
            # Load vectorstore
            vectorstore = self.load_vectorstore(document_id)
            if not vectorstore:
                return "Error: Could not load document data. Please ensure the document is properly processed."
            
            # Set default search_kwargs
            if search_kwargs is None:
                search_kwargs = {"k": 3}
            
            # Validate k value
            k = search_kwargs.get("k", 3)
            if k < 3 or k > 5:
                search_kwargs["k"] = 3
            
            # Create retriever
            retriever = vectorstore.as_retriever(
                search_type="similarity",
                search_kwargs=search_kwargs
            )
            
            # Format documents function
            def format_docs(retrieved_docs):
                return "\n\n".join([doc.page_content for doc in retrieved_docs])
            
            # Create chain
            chain = RunnableParallel({
                "context": retriever | RunnableLambda(format_docs),
                "question": RunnablePassthrough()
            }) | self.prompt | self.model | StrOutputParser()
            
            # Get answer
            result = chain.invoke(question)
            return result
            
        except Exception as e:
            return f"Error processing question: {str(e)}"